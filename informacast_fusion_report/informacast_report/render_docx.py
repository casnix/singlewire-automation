from __future__ import annotations

import datetime
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

from .crawler import FacilityReport, InstanceReport
from .resources import GROUPS

ACCENT = RGBColor(0x1F, 0x5F, 0x8B)
MUTED = RGBColor(0x5B, 0x66, 0x75)


def _stringify(item: dict, field: str) -> str:
    resolved = item.get(f"{field}_resolved")
    if resolved is not None:
        if isinstance(resolved, list):
            return ", ".join(str(v) for v in resolved) or "—"
        return str(resolved)

    value = item.get(field)
    if value is None:
        return "—"
    if isinstance(value, dict):
        return str(value.get("name", value))
    if isinstance(value, list):
        return ", ".join(str(v) for v in value) if value else "—"
    return str(value)


def render_docx(report: InstanceReport, output_path: str) -> None:
    doc = Document()

    title = doc.add_heading("InformaCast Fusion Configuration Report", level=0)
    subtitle = doc.add_paragraph("Full read-only export of instance configuration via the Fusion REST API")
    subtitle.runs[0].font.color.rgb = MUTED

    meta = doc.add_paragraph()
    meta.add_run(
        f"Generated {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')} · "
        f"API base {report.base_url} · "
        f"{len(report.facilities)} facilit{'y' if len(report.facilities) == 1 else 'ies'} crawled"
    ).font.size = Pt(9)

    for facility_report in report.facilities:
        _render_facility(doc, facility_report)

    doc.save(output_path)


def _render_facility(doc: Document, fr: FacilityReport) -> None:
    doc.add_page_break()
    facility_name = fr.facility["name"] if fr.facility else "Instance (no Facilities configured)"
    h = doc.add_heading(f"Facility: {facility_name}", level=1)
    if fr.facility:
        p = doc.add_paragraph(f"Facility ID: {fr.facility['id']}")
        p.runs[0].font.size = Pt(9)
        p.runs[0].font.color.rgb = MUTED

    for group_key, group_label in GROUPS.items():
        group_resources = [r for r in fr.resources.values() if r.spec.group == group_key]
        if not group_resources:
            continue

        doc.add_heading(group_label, level=2)

        for result in group_resources:
            singleton_tag = " [singleton]" if result.spec.is_singleton else ""
            heading = doc.add_heading(f"{result.spec.label} ({len(result.items)}){singleton_tag}", level=3)

            if result.spec.notes:
                note = doc.add_paragraph(result.spec.notes)
                note.runs[0].italic = True
                note.runs[0].font.size = Pt(9)

            if result.error:
                err = doc.add_paragraph(f"Not available: {result.error}")
                err.runs[0].font.color.rgb = RGBColor(0xB0, 0x30, 0x30)
                continue

            if not result.items:
                empty = doc.add_paragraph(f"No {result.spec.label.lower()} are configured.")
                empty.runs[0].italic = True
                continue

            _render_table(doc, result.items)

    _render_sites(doc, fr)
    _render_extensions(doc, fr)
    _render_alarms(doc, fr)


def _render_table(doc: Document, items: list[dict]) -> None:
    # Cap the number of columns shown for very wide/nested objects — the
    # HTML report is the better place to see every raw attribute; the Word
    # report favors readability over completeness for wide resources.
    sample = items[0]
    fields = [
        f for f in sample.keys()
        if not f.endswith("_resolved") and f not in ("permissions",)
    ]
    if len(fields) > 8:
        fields = fields[:8]

    table = doc.add_table(rows=1, cols=len(fields))
    table.style = "Light Grid Accent 1"
    hdr_cells = table.rows[0].cells
    for i, f in enumerate(fields):
        hdr_cells[i].text = f
        for run in hdr_cells[i].paragraphs[0].runs:
            run.bold = True

    for item in items:
        row_cells = table.add_row().cells
        for i, f in enumerate(fields):
            row_cells[i].text = _stringify(item, f)

    doc.add_paragraph()  # spacing after table


def _render_sites(doc: Document, fr: FacilityReport) -> None:
    doc.add_heading("Sites & Locations (Detail Tree)", level=2)
    if not fr.sites_tree:
        p = doc.add_paragraph("No sites configured (or not visible to this API account).")
        p.runs[0].italic = True
        return

    for site in fr.sites_tree:
        doc.add_heading(site.get("name", "Unnamed site"), level=4)
        for building in site.get("buildings", []):
            doc.add_paragraph(f"Building: {building.get('name', 'Unnamed')}", style="List Bullet")
            for floor in building.get("floors", []):
                doc.add_paragraph(f"Floor: {floor.get('name', 'Unnamed')}", style="List Bullet 2")
                for zone in floor.get("zones", []):
                    doc.add_paragraph(f"Zone: {zone.get('name', 'Unnamed')}", style="List Bullet 3")


def _render_extensions(doc: Document, fr: FacilityReport) -> None:
    doc.add_heading("Extensions Detail (Devices & Endpoints)", level=2)
    if not fr.extension_tree:
        p = doc.add_paragraph("No extensions configured (or not visible to this API account).")
        p.runs[0].italic = True
        return

    for ext in fr.extension_tree:
        name = ext.get("name") or ext.get("id", "Unnamed extension")
        disabled_tag = " (disabled)" if ext.get("disabled") else ""
        doc.add_heading(f"{name}{disabled_tag}", level=4)
        doc.add_paragraph(
            f"{len(ext.get('devices', []))} device(s), {len(ext.get('endpoints', []))} endpoint(s)"
        )
        for ep in ext.get("endpoints", []):
            ep_name = ep.get("name") or ep.get("id", "Unnamed endpoint")
            ep_type = f" ({ep['type']})" if ep.get("type") else ""
            doc.add_paragraph(f"Endpoint: {ep_name}{ep_type}", style="List Bullet")


def _render_alarms(doc: Document, fr: FacilityReport) -> None:
    doc.add_heading("Alarm Detail (Actions & Events)", level=2)
    if not fr.alarm_details:
        p = doc.add_paragraph("No alarm detail available.")
        p.runs[0].italic = True
        return

    for alarm in fr.alarm_details:
        doc.add_paragraph(
            f"{alarm.get('type')} — status: {alarm.get('status')}, muted: {alarm.get('muted')}, "
            f"{len(alarm.get('actions', []))} action(s), {len(alarm.get('events', []))} recent event(s)",
            style="List Bullet",
        )
