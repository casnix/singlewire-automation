"""Renders a NarrativeDoc (see narrative.py) as a Word document, following
the house style extracted from Claude_Word_Template.docx:

- Title: 26pt bold, #1F5C99
- Org/instance line: 16pt, #444444
- Subtitle: 12pt italic, #666666
- NOTE callouts: 10pt italic, #1F5C99
- Heading 1 / Heading 2: #2E74B5 (16pt / 13pt)
- Heading 3: #1F4D78 (12pt)
- Tables: bold white text on #1F5C99 header row, #F5F5F5/white zebra
  striping on data rows -- matching the template's Document Control /
  Revision History tables exactly.

Structure mirrors that template's conventions (title/classification block,
a Document-Control-style metadata table, an italic framing note, numbered
body sections) adapted for an auto-generated operational document rather
than a hand-authored design doc.
"""
from __future__ import annotations

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

from .narrative import NarrativeDoc, NarrativeSection, NarrativeTable

TITLE_COLOR = RGBColor(0x1F, 0x5C, 0x99)
ORG_LINE_COLOR = RGBColor(0x44, 0x44, 0x44)
SUBTITLE_COLOR = RGBColor(0x66, 0x66, 0x66)
NOTE_COLOR = RGBColor(0x1F, 0x5C, 0x99)
HEADING_12_COLOR = RGBColor(0x2E, 0x74, 0xB5)
HEADING_3_COLOR = RGBColor(0x1F, 0x4D, 0x78)
HEADER_FILL = "1F5C99"
ZEBRA_FILL = "F5F5F5"
WHITE = RGBColor(0xFF, 0xFF, 0xFF)


def render_narrative_docx(narrative: NarrativeDoc, output_path: str) -> None:
    doc = Document()

    _add_title_block(doc, narrative)
    _add_meta_table(doc, narrative.meta)
    _add_note(doc, narrative.intro_note)

    for section in narrative.sections:
        _add_section(doc, section)

    doc.save(output_path)


def _add_title_block(doc: Document, narrative: NarrativeDoc) -> None:
    p = doc.add_paragraph()
    run = p.add_run(narrative.title)
    run.bold = True
    run.font.size = Pt(26)
    run.font.color.rgb = TITLE_COLOR

    p = doc.add_paragraph()
    run = p.add_run(narrative.org_line)
    run.font.size = Pt(16)
    run.font.color.rgb = ORG_LINE_COLOR

    p = doc.add_paragraph()
    run = p.add_run(narrative.subtitle)
    run.italic = True
    run.font.size = Pt(12)
    run.font.color.rgb = SUBTITLE_COLOR

    doc.add_paragraph()


def _shade_cell(cell, fill_hex: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = tcPr.makeelement(qn("w:shd"), {})
        tcPr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)


def _add_meta_table(doc: Document, meta: dict) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = table.rows[0].cells
    hdr[0].text = "Field"
    hdr[1].text = "Value"
    for cell in hdr:
        _shade_cell(cell, HEADER_FILL)
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.color.rgb = WHITE
                r.font.size = Pt(10)

    for i, (key, value) in enumerate(meta.items()):
        row_cells = table.add_row().cells
        row_cells[0].text = key
        row_cells[1].text = str(value)
        fill = ZEBRA_FILL if i % 2 == 0 else "FFFFFF"
        for cell in row_cells:
            _shade_cell(cell, fill)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)

    doc.add_paragraph()


def _add_note(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(f"NOTE: {text}")
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = NOTE_COLOR
    doc.add_paragraph()


def _add_section(doc: Document, section: NarrativeSection) -> None:
    heading = doc.add_heading(section.heading, level=section.level)
    for run in heading.runs:
        if section.level == 3:
            run.font.color.rgb = HEADING_3_COLOR
        else:
            run.font.color.rgb = HEADING_12_COLOR

    for para_text in section.paragraphs:
        doc.add_paragraph(para_text)

    for note_text in section.notes:
        _add_note(doc, note_text)

    for table_model in section.tables:
        _add_narrative_table(doc, table_model)


def _add_narrative_table(doc: Document, table_model: NarrativeTable) -> None:
    table = doc.add_table(rows=1, cols=len(table_model.headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(table_model.headers):
        hdr_cells[i].text = str(header)
        _shade_cell(hdr_cells[i], HEADER_FILL)
        for p in hdr_cells[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.color.rgb = WHITE
                r.font.size = Pt(10)

    for row_i, row in enumerate(table_model.rows):
        row_cells = table.add_row().cells
        for col_i, value in enumerate(row):
            row_cells[col_i].text = str(value)
            for p in row_cells[col_i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9.5)
        fill = ZEBRA_FILL if row_i % 2 == 0 else "FFFFFF"
        for cell in row_cells:
            _shade_cell(cell, fill)

    if table_model.truncated_count:
        p = doc.add_paragraph()
        run = p.add_run(f"...and {table_model.truncated_count} more not shown here — see the JSON/HTML data report for the full list.")
        run.italic = True
        run.font.size = Pt(9)
        run.font.color.rgb = SUBTITLE_COLOR

    doc.add_paragraph()
