"""Not part of the shipped tool — a throwaway harness to validate that the
crawler + both renderers work end-to-end without needing a real Fusion
instance. Uses canned responses shaped like real API payloads.
"""
import json
import logging
import sys
sys.path.insert(0, ".")

from unittest.mock import patch, MagicMock

from informacast_report.api_client import ApiError, FusionApiClient, MAX_PAGES_PER_RESOURCE
from informacast_report.config import Settings
from informacast_report.crawler import Crawler
from informacast_report.logging_utils import setup_logging
from informacast_report.resources import RESOURCES

FAKE_DATA = {
    "/facilities": [],  # no facilities in this fake instance
    "/users": [
        {"id": "u1", "name": "Jane Admin", "email": "jane@example.com", "isLocked": False},
        {"id": "u2", "name": "API Service Account", "email": "svc@example.com", "isLocked": False},
    ],
    "/security-groups": [
        {"id": "sg1", "name": "Superuser", "userIds": ["u1"]},
    ],
    "/distribution-lists": [
        {"id": "dl1", "name": "All Staff", "createdAt": "2024-01-01T00:00:00Z"},
        {"id": "dl2", "name": "Front Desk", "createdAt": "2024-02-01T00:00:00Z"},
    ],
    "/message-templates": [
        {
            "id": "mt1", "name": "Severe Weather", "subject": "Severe weather alert",
            "body": "Take shelter now.", "distributionListIds": ["dl1"],
        },
    ],
    "/scenarios": [
        {"id": "sc1", "name": "Panic Button", "locationEnabled": True},
    ],
    "/sites": [
        {"id": "site1", "name": "Main Campus"},
    ],
    "/sites/site1/buildings": [
        {"id": "b1", "name": "West Building"},
    ],
    "/sites/site1/buildings/b1/floors": [
        {"id": "f1", "name": "1st Floor"},
    ],
    "/sites/site1/buildings/b1/floors/f1/zones": [
        {"id": "z1", "name": "Lobby"},
    ],
    "/alarms": [
        {"id": "al1", "type": "fusion_server_red", "status": "OK", "muted": False},
    ],
    "/alarms/al1/actions": [],
    "/alarms/al1/events": [],
}

# Singleton (non-list) resources are handled separately -- a plain dict per
# path, fetched via get_one() rather than paged_get().
FAKE_SINGLETON_DATA = {
    "/settings": {"orgName": "Example Org", "timezone": "America/Chicago"},
    # /ip-speaker-settings intentionally omitted -- exercises the 404 path.
}


def fake_get_one(self, path, facility_id=None):
    if path in FAKE_SINGLETON_DATA:
        return FAKE_SINGLETON_DATA[path]
    raise ApiError(f"404 Not Found for {path}", 404, path)


def fake_paged_get(self, path, facility_id=None, extra_params=None, limit=100, stats=None, pagination_style="offset"):
    data = FAKE_DATA.get(path)
    if data is None:
        raise ApiError(f"404 Not Found for {path}", 404, path)
    if stats is not None:
        stats.update(pages=1, items=len(data), advertised_total=len(data), envelope="paginated", pagination_style=pagination_style)
    yield from data


def main():
    settings = Settings(token="fake-token", base_url="https://api.icmobile.singlewire.com/api/v1", timeout=30)
    client = FusionApiClient(settings)

    with patch.object(FusionApiClient, "paged_get", fake_paged_get), \
         patch.object(FusionApiClient, "get_one", fake_get_one):
        crawler = Crawler(client, specs=RESOURCES)
        report = crawler.run()

    print(f"Crawled {len(report.facilities)} facility(ies)")
    for fr in report.facilities:
        for key, result in fr.resources.items():
            status = f"ERROR: {result.error}" if result.error else f"{len(result.items)} item(s)"
            print(f"  {key}: {status}")
        print(f"  sites_tree: {fr.sites_tree}")
        print(f"  alarm_details: {fr.alarm_details}")

    from informacast_report.render_html import render_html
    html = render_html(report)
    with open("/tmp/smoke_report.html", "w") as f:
        f.write(html)
    print(f"\nHTML report: {len(html)} chars written to /tmp/smoke_report.html")
    assert "Jane Admin" in html
    assert "Severe Weather" in html
    assert "All Staff" in html  # resolved from distributionListIds
    assert "Lobby" in html

    from informacast_report.render_docx import render_docx
    render_docx(report, "/tmp/smoke_report.docx")
    print("DOCX report written to /tmp/smoke_report.docx")

    print("\nSMOKE TEST PASSED (basic crawl + render)\n")


def test_logging_levels():
    """Confirm --verbose surfaces per-resource progress lines and --debug
    additionally surfaces raw HTTP/pagination detail, using the real
    paged_get pagination logic against a mocked HTTP layer (not a mocked
    paged_get) so we're actually exercising the pagination code path.
    """
    print("=" * 70)
    print("Testing --verbose output")
    print("=" * 70)
    setup_logging(verbose=True, debug=False)
    _run_two_page_fetch()

    print()
    print("=" * 70)
    print("Testing --debug output")
    print("=" * 70)
    setup_logging(verbose=False, debug=True)
    _run_two_page_fetch()


def _run_two_page_fetch():
    """Fetch a resource that spans two pages, using a real mocked HTTP
    session so pagination, not just the crawler, is exercised.
    """
    settings = Settings(token="fake-token", base_url="https://fake.example.com/api/v1", timeout=30)
    client = FusionApiClient(settings)

    page1 = MagicMock(status_code=200, ok=True, content=b"x" * 500)
    page1.json.return_value = {
        "total": 3, "partial": True, "previous": None, "next": "1",
        "data": [{"id": "a", "name": "A"}, {"id": "b", "name": "B"}],
    }
    page2 = MagicMock(status_code=200, ok=True, content=b"x" * 200)
    page2.json.return_value = {
        "total": 3, "partial": False, "previous": "0", "next": None,
        "data": [{"id": "c", "name": "C"}],
    }
    client.session.get = MagicMock(side_effect=[page1, page2])

    items = list(client.paged_get("/fake-resource"))
    assert [i["id"] for i in items] == ["a", "b", "c"], items


def test_pagination_loop_guard():
    """A server that never stops reporting partial=True should be caught by
    MAX_PAGES_PER_RESOURCE rather than looping forever.
    """
    print()
    print("=" * 70)
    print("Testing pagination loop guard (should raise ApiError, not hang)")
    print("=" * 70)
    setup_logging(verbose=False, debug=False)

    settings = Settings(token="fake-token", base_url="https://fake.example.com/api/v1", timeout=30)
    client = FusionApiClient(settings)

    def infinite_page(*args, **kwargs):
        resp = MagicMock(status_code=200, ok=True, content=b"x")
        resp.json.return_value = {
            "total": 999999, "partial": True, "previous": None, "next": "1",
            "data": [{"id": "loop", "name": "Loop"}],
        }
        return resp

    client.session.get = MagicMock(side_effect=infinite_page)

    try:
        import informacast_report.api_client as api_client_mod
        original = api_client_mod.MAX_PAGES_PER_RESOURCE
        api_client_mod.MAX_PAGES_PER_RESOURCE = 5
        try:
            list(client.paged_get("/looping-resource"))
            raise AssertionError("Expected ApiError from loop guard, got no error")
        except ApiError as exc:
            print(f"Caught expected ApiError: {exc}")
        finally:
            api_client_mod.MAX_PAGES_PER_RESOURCE = original
    finally:
        pass

    print("Loop guard test PASSED")


def test_unreliable_partial_next_flags():
    """Regression test for the reported bug: an endpoint that returns full
    pages of data but never sets `partial`/`next` correctly (e.g. always
    partial=False, next=None, even when more data exists). The OLD logic
    would stop after page 1 and silently truncate. The FIXED logic must
    keep paging because a full page came back, and additionally cross-check
    against `total`.
    """
    print()
    print("=" * 70)
    print("Testing pagination with UNRELIABLE partial/next flags (the reported bug)")
    print("=" * 70)
    setup_logging(verbose=False, debug=False)

    settings = Settings(token="fake-token", base_url="https://fake.example.com/api/v1", timeout=30)
    client = FusionApiClient(settings)

    # 250 total users, page size 100, server ALWAYS reports partial=False
    # and next=None regardless of whether more data remains -- exactly the
    # kind of unreliable envelope that caused silent truncation before.
    all_users = [{"id": f"u{i}", "name": f"User {i}"} for i in range(250)]

    def broken_envelope_page(url, params=None, headers=None, timeout=None):
        offset = params["offset"]
        limit = params["limit"]
        page_data = all_users[offset: offset + limit]
        resp = MagicMock(status_code=200, ok=True, content=b"x" * 100)
        resp.json.return_value = {
            "total": len(all_users),
            "partial": False,   # <-- always False, even mid-pagination (buggy server)
            "previous": None,
            "next": None,        # <-- always None, even mid-pagination (buggy server)
            "data": page_data,
        }
        return resp

    client.session.get = MagicMock(side_effect=broken_envelope_page)

    stats = {}
    items = list(client.paged_get("/users", limit=100, stats=stats))

    print(f"  Requested: 250 users across a server that never sets partial/next correctly")
    print(f"  Collected: {len(items)} items in {stats.get('pages')} page(s)")

    assert len(items) == 250, (
        f"BUG STILL PRESENT: only collected {len(items)}/250 items — pagination "
        f"stopped early because it trusted partial/next alone."
    )
    assert stats.get("pages") == 3, f"Expected 3 pages (100+100+50), got {stats.get('pages')}"
    assert [i["id"] for i in items] == [f"u{i}" for i in range(250)], "Items out of order or missing"

    print("  ✓ All 250 items collected across 3 pages despite unreliable partial/next flags.")
    print("Unreliable-envelope regression test PASSED")


def test_diagnostic_mode():
    """Exercise the --test resource diagnostic path directly."""
    print()
    print("=" * 70)
    print("Testing --test diagnostic mode against a mocked instance")
    print("=" * 70)

    from informacast_report.diagnostics import test_resource

    settings = Settings(token="fake-token", base_url="https://fake.example.com/api/v1", timeout=30)
    client = FusionApiClient(settings)

    call_log = []

    def router(url, params=None, headers=None, timeout=None):
        call_log.append(url)
        resp = MagicMock(status_code=200, ok=True, content=b"x" * 50)
        if url.endswith("/facilities"):
            resp.json.return_value = []  # no facilities in use
        elif url.endswith("/users"):
            offset = params.get("offset", 0)
            data = [{"id": "u1", "name": "Jane"}, {"id": "u2", "name": "Bob"}][offset: offset + params["limit"]]
            resp.json.return_value = {"total": 2, "partial": False, "next": None, "data": data}
        else:
            resp.status_code = 404
            resp.ok = False
        return resp

    client.session.get = MagicMock(side_effect=router)

    ok = test_resource(client, "users")
    assert ok is True, "Expected test_resource to report success for consistent data"

    # Also confirm an unknown key is handled gracefully rather than raising.
    ok_unknown = test_resource(client, "not_a_real_resource")
    assert ok_unknown is False

    print("Diagnostic mode test PASSED")


def test_total_overshoot_bug():
    """Regression test for the bug reported this time: an endpoint that
    always claims partial=True (and always returns a full page) no matter
    how far past its own declared `total` you go — exactly the pasted log
    (total=448 but still fetching full pages at offset=9000). The OLD OR-
    based logic would keep going because partial/full-page always won. The
    FIXED logic must treat `total` as a hard ceiling: stop (truncating the
    final page if needed) the moment total_yielded reaches it, regardless
    of what partial/next claim.
    """
    print()
    print("=" * 70)
    print("Testing pagination OVERSHOOTING total (this session's reported bug)")
    print("=" * 70)
    setup_logging(verbose=False, debug=False)

    settings = Settings(token="fake-token", base_url="https://fake.example.com/api/v1", timeout=30)
    client = FusionApiClient(settings)

    TOTAL = 448
    LIMIT = 100
    call_count = {"n": 0}

    def always_full_page(url, params=None, headers=None, timeout=None):
        # Server ALWAYS returns a full 100-item page and ALWAYS claims
        # partial=True / next=<something>, no matter the offset -- it does
        # not actually respect offset for real data past a point, it just
        # keeps manufacturing full pages. This is what caused the runaway
        # in the pasted log (still full pages at offset=9000, total=448).
        call_count["n"] += 1
        offset = params["offset"]
        data = [{"id": f"u{offset + i}", "name": f"User {offset + i}"} for i in range(LIMIT)]
        resp = MagicMock(status_code=200, ok=True, content=b"x" * 100)
        resp.json.return_value = {
            "total": TOTAL,
            "partial": True,     # <-- always true, forever, regardless of offset
            "previous": None,
            "next": str(offset + LIMIT),  # <-- always advances, never null
            "data": data,
        }
        return resp

    client.session.get = MagicMock(side_effect=always_full_page)

    stats = {}
    items = list(client.paged_get("/users", limit=LIMIT, stats=stats))

    print(f"  Server: always full pages, always partial=True, total={TOTAL}")
    print(f"  HTTP calls made: {call_count['n']}")
    print(f"  Items collected: {len(items)}  (pages: {stats.get('pages')})")
    print(f"  Truncated:       {stats.get('truncated')}")

    assert call_count["n"] == 5, (
        f"BUG STILL PRESENT: made {call_count['n']} HTTP calls instead of the expected 5 "
        f"(ceil(448/100)) — pagination is overshooting total again, same as the pasted log "
        f"running past offset=9000."
    )
    assert len(items) == TOTAL, f"Expected exactly {TOTAL} items, got {len(items)}"
    assert stats.get("truncated") is True, "Expected the final page to be flagged as truncated"
    # Confirm no duplicate/fabricated ids beyond the real total leaked through.
    assert items[-1]["id"] == f"u{TOTAL - 1}", f"Unexpected last item: {items[-1]}"

    print("  ✓ Stopped at exactly 5 pages / 448 items instead of running past total.")
    print("Total-overshoot regression test PASSED")


def test_stuck_page_repeat_bug():
    """Regression test for the bug reported this round: an endpoint that
    ignores `offset` entirely and just keeps re-serving page 1, forever.
    Same root cause as the device-groups cursor bug, but for offset-based
    pagination. Must raise immediately (not wait for MAX_PAGES) since the
    exact same ids repeating is an unambiguous signal.
    """
    print()
    print("=" * 70)
    print("Testing STUCK pagination: server ignores offset, always returns page 1")
    print("=" * 70)
    setup_logging(verbose=False, debug=False)

    settings = Settings(token="fake-token", base_url="https://fake.example.com/api/v1", timeout=30)
    client = FusionApiClient(settings)

    call_count = {"n": 0}

    def always_page_one(url, params=None, headers=None, timeout=None):
        call_count["n"] += 1
        resp = MagicMock(status_code=200, ok=True, content=b"x" * 100)
        resp.json.return_value = {
            "total": 448,  # server claims a large total...
            "partial": True,
            "previous": None,
            "next": "100",
            # ...but always returns the exact same 100 ids regardless of offset.
            "data": [{"id": f"u{i}", "name": f"User {i}"} for i in range(100)],
        }
        return resp

    client.session.get = MagicMock(side_effect=always_page_one)

    try:
        list(client.paged_get("/users"))
        raise AssertionError("Expected ApiError for stuck pagination, got none")
    except ApiError as exc:
        print(f"  Caught expected ApiError after {call_count['n']} call(s): {exc}")
        assert call_count["n"] == 2, (
            f"Expected exactly 2 HTTP calls (page 1, then page 2 detected as an exact "
            f"repeat) before aborting, got {call_count['n']} — it kept going instead of "
            f"failing fast."
        )
        assert "exact same" in str(exc) and "stuck" in str(exc)

    print("Stuck-page regression test PASSED")


def test_partial_duplicate_overlap():
    """A page that partially overlaps with previously-seen ids (rather than
    being an exact repeat) should be filtered/de-duplicated and logged as a
    warning, not silently accepted and not treated as fatal.
    """
    print()
    print("=" * 70)
    print("Testing PARTIAL duplicate overlap across pages (warn + dedupe, not fatal)")
    print("=" * 70)
    setup_logging(verbose=False, debug=False)

    settings = Settings(token="fake-token", base_url="https://fake.example.com/api/v1", timeout=30)
    client = FusionApiClient(settings)

    # Page 1: users 0-4. Page 2: overlaps on users 3-4, then adds 5-7 (new).
    # Page 3: empty -> stop. No `total` field at all on this endpoint.
    pages = [
        {"partial": True, "next": "1", "data": [{"id": f"u{i}"} for i in range(0, 5)]},
        {"partial": True, "next": "2", "data": [{"id": f"u{i}"} for i in range(3, 8)]},
        {"partial": False, "next": None, "data": []},
    ]
    call = {"n": 0}

    def router(url, params=None, headers=None, timeout=None):
        idx = min(call["n"], len(pages) - 1)
        call["n"] += 1
        resp = MagicMock(status_code=200, ok=True, content=b"x" * 50)
        resp.json.return_value = pages[idx]
        return resp

    client.session.get = MagicMock(side_effect=router)

    stats = {}
    items = list(client.paged_get("/users", limit=5, stats=stats))
    ids = [i["id"] for i in items]

    print(f"  Collected ids: {ids}")
    print(f"  Stats: {stats}")

    assert ids == [f"u{i}" for i in range(8)], f"Expected u0..u7 with no dupes, got {ids}"
    assert stats.get("duplicates") == 2, f"Expected 2 duplicates filtered (u3, u4 reappearing), got {stats.get('duplicates')}"
    assert stats.get("raw_items") == 10, f"Expected 10 raw items received (5+5), got {stats.get('raw_items')}"
    assert stats.get("items") == 8, f"Expected 8 unique items, got {stats.get('items')}"

    print("  ✓ Duplicates correctly filtered, unique count and raw count both correct.")
    print("Partial-overlap regression test PASSED")


def test_cursor_style_pagination():
    """Confirm pagination_style='cursor' actually works end-to-end: no
    offset param sent, `start` echoes back the prior `next` value, and it
    correctly stops when `next` goes null -- mirroring the standalone
    list_device_groups.py fix, but inside the main tool's generic client.
    """
    print()
    print("=" * 70)
    print("Testing pagination_style='cursor' (device-groups style)")
    print("=" * 70)
    setup_logging(verbose=False, debug=False)

    settings = Settings(token="fake-token", base_url="https://fake.example.com/api/v1", timeout=30)
    client = FusionApiClient(settings)

    pages = [
        {"total": 5, "next": "tokA", "data": [{"id": "g1"}, {"id": "g2"}]},
        {"total": 5, "next": "tokB", "data": [{"id": "g3"}, {"id": "g4"}]},
        {"total": 5, "next": None, "data": [{"id": "g5"}]},
    ]
    call_params = []

    def router(url, params=None, headers=None, timeout=None):
        call_params.append(dict(params))
        idx = len(call_params) - 1
        resp = MagicMock(status_code=200, ok=True, content=b"x" * 50)
        resp.json.return_value = pages[idx]
        return resp

    client.session.get = MagicMock(side_effect=router)

    items = list(client.paged_get("/device-groups", limit=2, pagination_style="cursor"))

    print(f"  Params sent per call: {call_params}")
    print(f"  Collected: {[i['id'] for i in items]}")

    assert [i["id"] for i in items] == ["g1", "g2", "g3", "g4", "g5"]
    # Crucially: no 'offset' key should ever be sent in cursor mode, and
    # 'start' should be absent on call 1, then exactly the prior 'next'.
    assert "offset" not in call_params[0]
    assert "start" not in call_params[0]
    assert call_params[1]["start"] == "tokA"
    assert call_params[2]["start"] == "tokB"

    print("  ✓ Cursor-style pagination correctly echoed 'next' back as 'start', no offset sent.")
    print("Cursor-style pagination test PASSED")


def test_device_groups_resource_is_cursor_style():
    """Guard against silently regressing the device_groups ResourceSpec back
    to offset-style pagination -- this was the concrete, confirmed fix.
    """
    print()
    print("=" * 70)
    print("Confirming resources.py: device_groups is configured as cursor-style")
    print("=" * 70)
    from informacast_report.resources import get_resource
    spec = get_resource("device_groups")
    assert spec.pagination_style == "cursor", (
        f"device_groups regressed to pagination_style={spec.pagination_style!r} -- "
        f"this endpoint is confirmed to need cursor-style pagination."
    )
    print("  ✓ device_groups is configured as pagination_style='cursor'")


def test_json_render_and_unit_filter():
    """Confirm render_json produces valid, complete JSON, and that
    resources_for_keys (--unit) correctly restricts to exact resource keys.
    """
    print()
    print("=" * 70)
    print("Testing JSON rendering and --unit resource-key filtering")
    print("=" * 70)

    from informacast_report.resources import RESOURCES, resources_for_keys
    from informacast_report.render_json import build_json_dict, render_json

    settings = Settings(token="fake-token", base_url="https://api.icmobile.singlewire.com/api/v1", timeout=30)
    client = FusionApiClient(settings)

    with patch.object(FusionApiClient, "paged_get", fake_paged_get):
        # --unit users,distribution_lists should crawl ONLY those two specs.
        unit_specs = resources_for_keys(["users", "distribution_lists"])
        assert [s.key for s in unit_specs] == ["users", "distribution_lists"]

        crawler = Crawler(client, specs=unit_specs)
        report = crawler.run()

    facility_resources = report.facilities[0].resources
    assert set(facility_resources.keys()) == {"users", "distribution_lists"}, (
        f"Expected only users+distribution_lists to be crawled with --unit, got "
        f"{set(facility_resources.keys())}"
    )

    # Unknown key must raise KeyError with a helpful message, not crash weirdly.
    try:
        resources_for_keys(["totally_bogus_key"])
        raise AssertionError("Expected KeyError for an unknown --unit key")
    except KeyError as exc:
        assert "Unknown resource key" in str(exc) and "Valid keys:" in str(exc)
        print(f"  ✓ Unknown --unit key correctly raises: {exc}")

    # Now render that restricted report as JSON and sanity-check structure.
    json_str = render_json(report)
    parsed = json.loads(json_str)  # must be valid JSON

    assert parsed["base_url"] == settings.base_url
    assert len(parsed["facilities"]) == 1
    resources_json = parsed["facilities"][0]["resources"]
    assert set(resources_json.keys()) == {"users", "distribution_lists"}
    assert resources_json["users"]["item_count"] == 2
    assert resources_json["users"]["items"][0]["name"] == "Jane Admin"
    assert resources_json["distribution_lists"]["pagination_style"] == "cursor"

    print(f"  ✓ JSON parses cleanly and contains exactly the --unit-filtered resources.")
    print(f"  Sample: {json_str[:200]}...")
    print("JSON render / --unit filter test PASSED")


def test_default_pagination_style_is_cursor():
    """Guard the core change from this round: every resource defaults to
    cursor-style pagination now, confirmed against the real API rather than
    the original 'offset' assumption.
    """
    print()
    print("=" * 70)
    print("Confirming resources.py: cursor is now the default for ALL resources")
    print("=" * 70)
    from informacast_report.resources import RESOURCES

    non_cursor = [r.key for r in RESOURCES if r.pagination_style != "cursor"]
    assert not non_cursor, (
        f"Expected every resource to default to pagination_style='cursor', but these "
        f"don't: {non_cursor}"
    )
    print(f"  ✓ All {len(RESOURCES)} resources use pagination_style='cursor'.")


def test_global_pagination_style_override_via_dataclasses_replace():
    """Confirm the mechanism main.py uses to apply --pagination-style across
    an entire (non --test) crawl -- dataclasses.replace on every spec --
    actually produces specs the crawler will honor.
    """
    print()
    print("=" * 70)
    print("Testing global --pagination-style override (full crawl, not --test)")
    print("=" * 70)
    import dataclasses
    from informacast_report.resources import resources_for_keys

    specs = resources_for_keys(["users", "device_groups"])
    assert specs[0].pagination_style == "cursor" and specs[1].pagination_style == "cursor"

    overridden = [dataclasses.replace(s, pagination_style="offset") for s in specs]
    assert all(s.pagination_style == "offset" for s in overridden), (
        "Global override via dataclasses.replace did not apply to every spec"
    )
    # Original specs must be untouched (frozen dataclass, new instances only).
    assert specs[0].pagination_style == "cursor", "Original ResourceSpec was mutated!"

    print("  ✓ dataclasses.replace cleanly overrides style on copies without mutating originals.")


def test_new_granular_resources_registered():
    """Guard the additions from this round: DialCast, CallAware, Inbound
    CAP/Email/RSS, and finer-grained recipient/device resources should all
    be present and reachable by key, using paths CONFIRMED against the real
    OpenAPI spec (not the guesses from the previous round -- several of
    those were wrong and have been corrected/replaced).
    """
    print()
    print("=" * 70)
    print("Confirming newly added granular notification/device/recipient resources")
    print("=" * 70)
    from informacast_report.resources import get_resource

    expected_keys = [
        "dial_cast", "dial_cast_phone_exceptions", "inbound_cap_rules", "inbound_email",
        "inbound_rss_feeds", "active_callaware_calls", "gateways",
        "ip_speakers", "ip_speaker_sip_parameters", "ip_speaker_jobs", "ip_speaker_settings",
        "tts_voices", "tts_lexicons", "tts_defaults", "incidents", "facilities",
    ]
    for key in expected_keys:
        spec = get_resource(key)  # raises KeyError if missing
        if not spec.is_singleton:
            assert spec.pagination_style == "cursor"
        assert spec.path.startswith("/"), f"{key} has a suspicious path: {spec.path!r}"

    # Guard against the wrong guesses from last round silently coming back.
    from informacast_report.resources import RESOURCE_BY_KEY
    removed_wrong_guesses = [
        "inbound_cap", "inbound_rss", "call_aware_redirects", "paging_gateways",
        "roll_call", "endpoints", "recipient_group_tags", "desktop_notifiers", "domains",
    ]
    for key in removed_wrong_guesses:
        assert key not in RESOURCE_BY_KEY, (
            f"{key!r} should have been removed/replaced after spec validation, but it's back"
        )

    print(f"  ✓ All {len(expected_keys)} new resources are registered and reachable.")
    print(f"  ✓ All {len(removed_wrong_guesses)} incorrect prior guesses are confirmed gone.")


def test_singleton_resource_handling():
    """Confirm is_singleton resources (settings, ip_speaker_settings) are
    fetched via get_one() and wrapped as a 1-item list, not treated as a
    paginated list (which would have silently returned zero items, as it
    did before this was fixed).
    """
    print()
    print("=" * 70)
    print("Testing singleton resource handling (e.g. /settings)")
    print("=" * 70)
    from informacast_report.resources import get_resource, resources_for_keys

    settings_spec = get_resource("settings")
    assert settings_spec.is_singleton is True

    settings = Settings(token="fake-token", base_url="https://fake.example.com/api/v1", timeout=30)
    client = FusionApiClient(settings)
    client.get_one = MagicMock(return_value={"orgName": "Example Org"})

    crawler = Crawler(client, specs=resources_for_keys(["settings"]))
    # Bypass /facilities lookup for this narrow test -- just exercise one facility context.
    fr = crawler._crawl_facility(None)

    result = fr.resources["settings"]
    assert result.error is None
    assert len(result.items) == 1
    assert result.items[0] == {"orgName": "Example Org"}
    assert result.pagination_stats.get("envelope") == "singleton"
    client.get_one.assert_called_once()

    print("  ✓ Singleton resource fetched via get_one() and wrapped correctly, not silently empty.")


def test_extension_tree_crawl():
    """Confirm the new nested Extension -> Devices/Endpoints crawl works,
    since Endpoints don't exist as a flat top-level resource in the real API
    (confirmed via OpenAPI spec) -- they're only reachable per-Extension.
    """
    print()
    print("=" * 70)
    print("Testing Extension tree crawl (nested Devices & Endpoints)")
    print("=" * 70)
    from informacast_report.resources import resources_for_keys

    settings = Settings(token="fake-token", base_url="https://fake.example.com/api/v1", timeout=30)
    client = FusionApiClient(settings)

    extensions_data = [{"id": "ext1", "name": "SchoolMessenger"}]
    nested_data = {
        "/extensions/ext1/devices": [{"id": "dev1", "name": "Device A"}],
        "/extensions/ext1/endpoints": [{"id": "ep1", "name": "Endpoint A", "type": "SCHOOL_MESSENGER"}],
    }

    def fake_paged(self, path, facility_id=None, extra_params=None, limit=100, stats=None, pagination_style="cursor"):
        if path == "/extensions":
            yield from extensions_data
        elif path in nested_data:
            yield from nested_data[path]
        else:
            raise ApiError(f"404 Not Found for {path}", 404, path)

    with patch.object(FusionApiClient, "paged_get", fake_paged):
        crawler = Crawler(client, specs=resources_for_keys(["extensions"]))
        fr = crawler._crawl_facility(None)

    assert len(fr.extension_tree) == 1
    ext = fr.extension_tree[0]
    assert ext["name"] == "SchoolMessenger"
    assert ext["devices"] == [{"id": "dev1", "name": "Device A"}]
    assert ext["endpoints"] == [{"id": "ep1", "name": "Endpoint A", "type": "SCHOOL_MESSENGER"}]

    print("  ✓ Extension tree correctly nests Devices and Endpoints per-extension.")


def test_facility_terminology_and_header():
    """Confirm the Facility fix: /facilities is called (not /domains), and
    the x-singlewire-facility header (not x-singlewire-domain) is sent.
    """
    print()
    print("=" * 70)
    print("Testing Facility terminology/header fix")
    print("=" * 70)
    from informacast_report.crawler import list_facilities

    settings = Settings(token="fake-token", base_url="https://fake.example.com/api/v1", timeout=30)
    client = FusionApiClient(settings)

    calls = []

    def fake_get(url, params=None, headers=None, timeout=None):
        calls.append((url, dict(headers or {})))
        resp = MagicMock(status_code=200, ok=True, content=b"x" * 50)
        if url.endswith("/facilities"):
            resp.json.return_value = {
                "total": 1, "partial": False, "next": None,
                "data": [{"id": "fac1", "name": "Main Campus"}],
            }
        else:
            resp.json.return_value = {"total": 0, "partial": False, "next": None, "data": []}
        return resp

    client.session.get = MagicMock(side_effect=fake_get)

    facilities = list_facilities(client)
    assert facilities == [{"id": "fac1", "name": "Main Campus"}]
    assert any(url.endswith("/facilities") for url, _ in calls), "list_facilities() did not call /facilities"
    assert not any(url.endswith("/domains") for url, _ in calls), "list_facilities() should never call /domains"

    # Now confirm the header name used when a facility_id is passed through.
    list(client.paged_get("/users", facility_id="fac1"))
    header_calls = [h for _, h in calls if h]
    assert any("x-singlewire-facility" in h for h in header_calls), (
        "Expected x-singlewire-facility header to be sent, got: " + str(header_calls)
    )
    assert not any("x-singlewire-domain" in h for h in header_calls), (
        "x-singlewire-domain should never be sent -- that header doesn't exist in the real API"
    )

    print("  ✓ /facilities called (not /domains), x-singlewire-facility header sent (not x-singlewire-domain).")


def test_narrative_end_to_end():
    """End-to-end test of the new instance-specific narrative feature:
    builds a realistic mocked instance (DialCast configs referencing both a
    shared Message Template and an embedded/direct recipient, an Extension
    with devices/endpoints, an empty Device Group, an unreferenced
    Distribution List, an unhealthy alarm) and confirms the narrative
    correctly resolves names, traces the DialCast -> notification ->
    recipients chain, and flags the anomalies -- then renders it to an
    actual .docx and verifies the file opens and contains expected text.
    """
    print()
    print("=" * 70)
    print("Testing instance-specific narrative generation end-to-end")
    print("=" * 70)

    from informacast_report.narrative import build_narrative
    from informacast_report.render_narrative_docx import render_narrative_docx
    from informacast_report.resources import RESOURCES

    settings = Settings(token="fake-token", base_url="https://api.icmobile.singlewire.com/api/v1", timeout=30)
    client = FusionApiClient(settings)

    fake_data = {
        "/facilities": [],
        "/users": [],
        "/distribution-lists": [
            {"id": "dl1", "name": "All Staff"},          # referenced by a template
            {"id": "dl2", "name": "Unused Overflow List"},  # NOT referenced -- should be flagged
        ],
        "/device-groups": [
            {"id": "dg1", "name": "Building A Speakers", "additionIds": ["p1", "p2"],
             "numPhones": 0, "numSpeakers": 12, "numIdns": 0},
            {"id": "dg2", "name": "Empty Leftover Group",  # should be flagged as empty
             "additionIds": [], "filters": None, "logicalExpression": None, "baseDeviceGroupIds": []},
        ],
        "/message-templates": [
            {"id": "mt1", "name": "Fire Alarm", "distributionListIds": ["dl1"],
             "deviceGroupIds": ["dg1"], "confirmationRequestId": None, "notificationProfileId": None,
             "incidentPlanId": None, "ttsVoiceId": None},
        ],
        "/dialcast-dialing-configurations": [
            {
                "id": "dc1", "name": "Fire Line", "dialingPatternRegex": "^5551\\d{3}$",
                "endpointIds": ["ep1"], "fallbackNotification": None,
                "notification": {"messageTemplateId": "mt1"},  # references shared template
            },
            {
                "id": "dc2", "name": "Direct Page Line", "dialingPatternRegex": "^5559999$",
                "endpointIds": [], "fallbackNotification": {"foo": "bar"},
                "notification": {"distributionListIds": ["dl1"], "deviceGroupIds": []},  # embedded, no template
            },
        ],
        "/dialcast-dialing-configurations/dc1/rule-actions": [
            {"id": "ra1", "name": "Notify Security Webhook"},
        ],
        "/dialcast-dialing-configurations/dc2/rule-actions": [],
        "/dialcast-phone-exceptions": [],
        "/extensions": [
            {"id": "ext1", "name": "SchoolMessenger Connector"},
        ],
        "/extensions/ext1/devices": [{"id": "d1", "name": "SM Device 1"}],
        "/extensions/ext1/endpoints": [{"id": "ep1", "name": "Fire Alarm Endpoint", "type": "SCHOOL_MESSENGER"}],
        "/scenarios": [],
        "/alarms": [
            {"id": "al1", "type": "fusion_server_red", "status": "CRITICAL", "muted": False},
            {"id": "al2", "type": "license_expiring", "status": "OK", "muted": False},
        ],
        "/alarms/al1/actions": [],
        "/alarms/al1/events": [],
        "/alarms/al2/actions": [],
        "/alarms/al2/events": [],
    }

    def fake_paged(self, path, facility_id=None, extra_params=None, limit=100, stats=None, pagination_style="cursor"):
        data = fake_data.get(path)
        if data is None:
            raise ApiError(f"404 Not Found for {path}", 404, path)
        yield from data

    with patch.object(FusionApiClient, "paged_get", fake_paged):
        specs = [
            r for r in RESOURCES
            if r.key in (
                "facilities", "users", "distribution_lists", "device_groups", "message_templates",
                "dial_cast", "dial_cast_phone_exceptions", "extensions", "scenarios", "alarms",
            )
        ]
        crawler = Crawler(client, specs=specs)
        report = crawler.run()

    narrative = build_narrative(report)
    fr = report.facilities[0]

    # -- Confirm the DialCast -> template chain resolved correctly --
    dial_cast_section = next(s for s in narrative.sections if s.heading == "DialCast: Dial-Pattern Triggers")
    table = dial_cast_section.tables[0]
    by_name = {row[0]: row for row in table.rows}

    assert "template: Fire Alarm" in by_name["Fire Line"][3], f"Expected Fire Line to resolve to its template, got: {by_name['Fire Line']}"
    assert by_name["Fire Line"][2] == "Fire Alarm Endpoint", "Endpoint id should resolve to its name via the Extension tree"
    assert by_name["Fire Line"][4] == "1", "Fire Line should show 1 linked Rule Action"

    assert "embedded notification" in by_name["Direct Page Line"][3]
    assert "All Staff" in by_name["Direct Page Line"][3], "Direct Page Line's embedded recipients should resolve to a name"
    assert by_name["Direct Page Line"][5] == "yes", "Direct Page Line has a fallback notification"

    print("  ✓ DialCast -> Message Template chain resolved correctly (Fire Line -> Fire Alarm template)")
    print("  ✓ DialCast -> embedded recipients resolved correctly (Direct Page Line -> All Staff)")
    print("  ✓ DialCast -> Extension endpoint name resolved correctly (ep1 -> Fire Alarm Endpoint)")
    print("  ✓ DialCast -> nested Rule Action count correct (Fire Line: 1, Direct Page Line: 0)")

    # -- Confirm anomaly detection fired correctly --
    anomalies_section = next(s for s in narrative.sections if s.heading == "Things Worth Verifying in This Facility")
    notes_text = " ".join(anomalies_section.notes)
    assert "Empty Leftover Group" in notes_text, "Empty device group should be flagged"
    assert "Unused Overflow List" in notes_text, "Unreferenced distribution list should be flagged"
    print("  ✓ Anomaly detection correctly flagged the empty Device Group and unreferenced Distribution List")

    # -- Confirm monitoring section flags the unhealthy, unmuted alarm --
    monitoring_section = next(s for s in narrative.sections if s.heading == "Monitoring & Operational State")
    monitoring_notes = " ".join(monitoring_section.notes)
    assert "1 alarm" in monitoring_notes
    print("  ✓ Unhealthy/unmuted alarm correctly flagged in Monitoring section")

    # -- Render to an actual docx and confirm it opens and contains expected content --
    out_path = "/tmp/narrative_test.docx"
    render_narrative_docx(narrative, out_path)

    from docx import Document as DocxDocument
    doc = DocxDocument(out_path)
    full_text = "\n".join(p.text for p in doc.paragraphs)
    for t in doc.tables:
        for row in t.rows:
            full_text += "\n" + " | ".join(c.text for c in row.cells)

    assert "InformaCast Fusion — Operational Narrative" in full_text
    assert "Fire Line" in full_text and "Fire Alarm" in full_text
    assert "Empty Leftover Group" in full_text
    print(f"  ✓ Rendered .docx opens correctly and contains expected content ({len(doc.paragraphs)} paragraphs, {len(doc.tables)} tables)")

    print("Narrative end-to-end test PASSED")


if __name__ == "__main__":
    main()
    test_logging_levels()
    test_pagination_loop_guard()
    test_unreliable_partial_next_flags()
    test_total_overshoot_bug()
    test_stuck_page_repeat_bug()
    test_partial_duplicate_overlap()
    test_cursor_style_pagination()
    test_device_groups_resource_is_cursor_style()
    test_default_pagination_style_is_cursor()
    test_global_pagination_style_override_via_dataclasses_replace()
    test_new_granular_resources_registered()
    test_singleton_resource_handling()
    test_extension_tree_crawl()
    test_facility_terminology_and_header()
    test_diagnostic_mode()
    test_json_render_and_unit_filter()
    test_narrative_end_to_end()
    print("\nALL SMOKE TESTS PASSED")
